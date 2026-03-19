"""Tests for export modules: PRISMA, CSV, Excel, DOCX, methods section."""

import csv
import json
import os
from pathlib import Path
from unittest.mock import patch

import openpyxl
import pytest

from engine.core.database import ReviewDatabase
from engine.core.review_spec import load_review_spec
from engine.exporters import export_all
from engine.exporters.docx_export import export_evidence_docx
from engine.exporters.evidence_table import (
    NO_EXTRACTION_MARKER,
    export_evidence_csv,
    export_evidence_excel,
)
from engine.exporters.methods_section import generate_methods_section, export_methods_md
from engine.exporters.prisma import generate_prisma_flow, export_prisma_csv
from engine.search.models import Citation

SPEC_PATH = Path(__file__).resolve().parent.parent / "review_specs" / "surgical_autonomy_v1.yaml"


@pytest.fixture(scope="module")
def spec():
    return load_review_spec(SPEC_PATH)


@pytest.fixture()
def populated_db(tmp_path, spec):
    """Create a DB with papers at various pipeline stages and extraction data."""
    db = ReviewDatabase("test_export", data_root=tmp_path)

    # 10 PubMed + 5 OpenAlex papers
    pm_cits = [
        Citation(title=f"PubMed Study {i}", source="pubmed", pmid=str(i),
                 doi=f"10.1/{i}", authors=["Smith A", "Jones B"],
                 journal="J Surg Robot", year=2023)
        for i in range(1, 11)
    ]
    oa_cits = [
        Citation(title=f"OpenAlex Study {i}", source="openalex", pmid=str(100 + i),
                 doi=f"10.2/{i}", authors=["Lee C"], journal="Robot Rev", year=2024)
        for i in range(1, 6)
    ]
    db.add_papers(pm_cits + oa_cits)

    papers = db.get_papers_by_status("INGESTED")

    # Screen 8 in, 4 out, 3 flagged
    for p in papers[:8]:
        db.add_screening_decision(p["id"], 1, "include", "Relevant", "qwen3:8b")
        db.add_screening_decision(p["id"], 2, "include", "Confirmed", "qwen3:8b")
        db.update_status(p["id"], "ABSTRACT_SCREENED_IN")

    for p in papers[8:12]:
        db.add_screening_decision(p["id"], 1, "exclude", "Not surgical", "qwen3:8b")
        db.add_screening_decision(p["id"], 2, "exclude", "Confirmed exclude", "qwen3:8b")
        db.update_status(p["id"], "ABSTRACT_SCREENED_OUT")

    for p in papers[12:15]:
        db.add_screening_decision(p["id"], 1, "include", "Maybe relevant", "qwen3:8b")
        db.add_screening_decision(p["id"], 2, "exclude", "Borderline", "qwen3:8b")
        db.update_status(p["id"], "ABSTRACT_SCREEN_FLAGGED")

    # Walk 5 screened-in papers to EXTRACTED/AI_AUDIT_COMPLETE
    screened_in = db.get_papers_by_status("ABSTRACT_SCREENED_IN")
    schema_hash = spec.extraction_hash()

    for j, p in enumerate(screened_in[:5]):
        pid = p["id"]
        db.update_status(pid, "PDF_ACQUIRED")
        db.update_status(pid, "PARSED")
        db.update_status(pid, "EXTRACTED")

        ext_id = db.add_extraction(
            pid, schema_hash,
            {"study_design": "RCT", "sample_size": "20"},
            "reasoning trace here", "deepseek-r1:32b",
        )

        # Add evidence spans for a few fields
        db.add_evidence_span(ext_id, "study_design", "RCT", "An RCT was performed.", 0.95)
        db.add_evidence_span(ext_id, "sample_size", "20", "Twenty trials.", 0.9)
        db.add_evidence_span(ext_id, "robot_platform", "STAR", "The STAR robot.", 0.85)

        # Audit 3 papers
        if j < 3:
            spans = db._conn.execute(
                "SELECT id FROM evidence_spans WHERE extraction_id = ?", (ext_id,)
            ).fetchall()
            for s in spans:
                db.update_audit(s["id"], "verified", "qwen3:32b", "Confirmed.")
            db.update_status(pid, "AI_AUDIT_COMPLETE")

    yield db
    db.close()


# ── PRISMA Flow ──────────────────────────────────────────────────────


def test_prisma_flow_counts(populated_db):
    flow = generate_prisma_flow(populated_db)
    assert flow["records_identified"] == 15
    assert flow["records_by_source"]["pubmed"] == 10
    assert flow["records_by_source"]["openalex"] == 5
    assert flow["records_excluded"] == 4
    assert flow["screen_flagged"] == 3
    assert flow["studies_included"] == 3  # AI_AUDIT_COMPLETE


def test_prisma_csv(populated_db, tmp_path):
    out = str(tmp_path / "prisma.csv")
    export_prisma_csv(populated_db, out)
    assert Path(out).exists()

    with open(out) as f:
        reader = csv.reader(f)
        rows = list(reader)
    # Header + data rows
    assert len(rows) > 5
    assert rows[0] == ["Stage", "Count", "Detail"]


# ── Evidence CSV ─────────────────────────────────────────────────────


def test_evidence_csv_columns(populated_db, spec, tmp_path):
    out = str(tmp_path / "evidence.csv")
    export_evidence_csv(populated_db, spec, out)
    assert Path(out).exists()

    with open(out) as f:
        reader = csv.reader(f)
        rows = list(reader)

    headers = rows[0]
    # Base columns present
    assert "paper_id" in headers
    assert "pmid" in headers
    assert "title" in headers

    # Extraction field columns present
    assert "study_design" in headers
    assert "study_design_snippet" in headers
    assert "study_design_confidence" in headers
    assert "study_design_audit" in headers

    # 3 papers at AI_AUDIT_COMPLETE (exporters no longer include EXTRACTED)
    assert len(rows) - 1 == 3


# ── Evidence Excel ───────────────────────────────────────────────────


def test_evidence_excel_sheets(populated_db, spec, tmp_path):
    out = str(tmp_path / "evidence.xlsx")
    export_evidence_excel(populated_db, spec, out)
    assert Path(out).exists()

    wb = openpyxl.load_workbook(out)
    assert wb.sheetnames == ["Evidence Table", "Screening Log", "Audit Log"]

    # Evidence Table has header + data rows
    ws1 = wb["Evidence Table"]
    assert ws1.max_row >= 4  # 1 header + 3 AI_AUDIT_COMPLETE papers

    # Screening Log has entries
    ws2 = wb["Screening Log"]
    assert ws2.max_row > 1

    # Audit Log has entries
    ws3 = wb["Audit Log"]
    assert ws3.max_row > 1
    wb.close()


# ── DOCX ─────────────────────────────────────────────────────────────


def test_docx_created(populated_db, spec, tmp_path):
    out = str(tmp_path / "evidence.docx")
    export_evidence_docx(populated_db, spec, out)
    assert Path(out).exists()

    # Verify it's a valid docx by loading it
    from docx import Document
    doc = Document(out)
    # Should have at least one table
    assert len(doc.tables) >= 1
    # Table should have header + data rows
    table = doc.tables[0]
    assert len(table.rows) >= 4  # 1 header + 3 AI_AUDIT_COMPLETE papers


# ── Methods Section ──────────────────────────────────────────────────


def test_methods_section_content(populated_db, spec):
    methods = generate_methods_section(populated_db, spec)

    # Key pipeline details present
    assert "PubMed" in methods
    assert "OpenAlex" in methods
    # Model names should come from spec, not be hardcoded
    assert spec.screening_models.primary in methods
    assert "deepseek-r1:32b" in methods  # from DB extraction data
    assert "qwen3:32b" in methods  # from DB audit data
    assert "dual-pass" in methods
    assert "two-pass" in methods
    assert "15" in methods  # total records


def test_methods_md_export(populated_db, spec, tmp_path):
    out = str(tmp_path / "methods.md")
    export_methods_md(populated_db, spec, out)
    assert Path(out).exists()

    content = Path(out).read_text()
    assert content.startswith("# Methods")
    assert "systematic search" in content


# ── export_all ───────────────────────────────────────────────────────


def test_export_all(populated_db, spec, tmp_path):
    out_dir = str(tmp_path / "all_exports")
    paths = export_all(populated_db, spec, "test_export", output_dir=out_dir)

    expected_keys = {
        "prisma_csv", "evidence_csv", "evidence_xlsx", "evidence_docx", "methods_md",
        "trace_quality_report", "trace_quality_report_md", "traces_dir",
    }
    assert expected_keys.issubset(set(paths.keys()))

    for key, path in paths.items():
        assert Path(path).exists(), f"{key} not found at {path}"


# ── H6: Atomic write — no partial files on error ────────────────────


def test_atomic_csv_no_partial_on_error(populated_db, spec, tmp_path):
    """If CSV export fails mid-write, no final file or temp file remains."""
    out = str(tmp_path / "evidence.csv")

    with patch("engine.exporters.evidence_table.csv.writer") as mock_writer:
        # Let writerow succeed for header, fail on writerows
        instance = mock_writer.return_value
        instance.writerows.side_effect = IOError("disk full")

        with pytest.raises(IOError, match="disk full"):
            export_evidence_csv(populated_db, spec, out)

    assert not Path(out).exists(), "Final file should not exist after error"
    assert not Path(out + ".tmp").exists(), "Temp file should be cleaned up"


def test_atomic_docx_no_partial_on_error(populated_db, spec, tmp_path):
    """If DOCX export fails during save, no final file or temp file remains."""
    out = str(tmp_path / "evidence.docx")

    with patch("engine.exporters.docx_export.Document") as mock_doc_cls:
        mock_doc = mock_doc_cls.return_value
        mock_doc.sections = [type("Sec", (), {
            "orientation": None, "page_width": 1, "page_height": 2,
            "left_margin": 1, "right_margin": 1, "top_margin": 1, "bottom_margin": 1,
        })()]
        mock_doc.add_paragraph.return_value = type("Para", (), {
            "add_run": lambda self, *a, **kw: type("Run", (), {
                "bold": False, "font": type("F", (), {"size": None})()
            })()
        })()
        mock_doc.add_table.return_value = type("T", (), {
            "style": None,
            "rows": [type("R", (), {"cells": [type("C", (), {
                "text": "", "paragraphs": []
            })() for _ in range(50)]})() for _ in range(20)]
        })()
        mock_doc.save.side_effect = IOError("disk full")

        with pytest.raises(IOError, match="disk full"):
            export_evidence_docx(populated_db, spec, out)

    assert not Path(out).exists(), "Final file should not exist after error"
    assert not Path(out + ".tmp").exists(), "Temp file should be cleaned up"


def test_atomic_prisma_csv_no_partial_on_error(populated_db, tmp_path):
    """If PRISMA CSV export fails, no final or temp file remains."""
    out = str(tmp_path / "prisma.csv")

    with patch("engine.exporters.prisma.csv.writer") as mock_writer:
        instance = mock_writer.return_value
        instance.writerows.side_effect = IOError("disk full")

        with pytest.raises(IOError, match="disk full"):
            export_prisma_csv(populated_db, out)

    assert not Path(out).exists()
    assert not Path(out + ".tmp").exists()


def test_atomic_methods_md_no_partial_on_error(populated_db, spec, tmp_path):
    """If methods MD export fails during write, no file remains."""
    out = str(tmp_path / "methods.md")

    with patch("builtins.open", side_effect=IOError("disk full")):
        with pytest.raises(IOError, match="disk full"):
            export_methods_md(populated_db, spec, out)

    assert not Path(out).exists()


# ── H13: Empty extraction rows marked ───────────────────────────────


@pytest.fixture()
def db_with_empty_extractions(tmp_path, spec):
    """DB with one paper that has extractions and one that doesn't."""
    db = ReviewDatabase("test_empty", data_root=tmp_path)

    # Two papers
    cits = [
        Citation(title=f"Study {i}", source="pubmed", pmid=str(i),
                 doi=f"10.1/{i}", authors=["Auth A"], journal="J Test", year=2023)
        for i in range(1, 3)
    ]
    db.add_papers(cits)
    papers = db.get_papers_by_status("INGESTED")
    schema_hash = spec.extraction_hash()

    for p in papers:
        db.add_screening_decision(p["id"], 1, "include", "Relevant", "qwen3:8b")
        db.add_screening_decision(p["id"], 2, "include", "Confirmed", "qwen3:8b")
        db.update_status(p["id"], "ABSTRACT_SCREENED_IN")
        db.update_status(p["id"], "PDF_ACQUIRED")
        db.update_status(p["id"], "PARSED")
        db.update_status(p["id"], "EXTRACTED")

    # Paper 1: has extraction data
    p1 = papers[0]["id"]
    ext_id = db.add_extraction(
        p1, schema_hash,
        {"study_design": "RCT"}, "trace", "deepseek-r1:32b",
    )
    db.add_evidence_span(ext_id, "study_design", "RCT", "An RCT.", 0.95)
    spans = db._conn.execute(
        "SELECT id FROM evidence_spans WHERE extraction_id = ?", (ext_id,)
    ).fetchall()
    for s in spans:
        db.update_audit(s["id"], "verified", "gemma3:27b", "OK")
    db.update_status(p1, "AI_AUDIT_COMPLETE")

    # Paper 2: has extraction row but NO evidence spans
    p2 = papers[1]["id"]
    db.add_extraction(
        p2, schema_hash, {}, "empty trace", "deepseek-r1:32b",
    )
    db.update_status(p2, "AI_AUDIT_COMPLETE")

    yield db
    db.close()


def test_empty_extraction_has_marker(db_with_empty_extractions, spec, tmp_path):
    """Papers with no extraction spans get [NO EXTRACTION DATA] marker."""
    out = str(tmp_path / "evidence.csv")
    export_evidence_csv(db_with_empty_extractions, spec, out)

    with open(out) as f:
        reader = csv.DictReader(f)
        rows = list(reader)

    assert len(rows) == 2

    # Find the empty paper (paper 2, no spans)
    markers = [r for r in rows if NO_EXTRACTION_MARKER in r.values()]
    assert len(markers) == 1, "Exactly one paper should have the marker"

    # The populated paper should NOT have the marker
    non_markers = [r for r in rows if NO_EXTRACTION_MARKER not in r.values()]
    assert len(non_markers) == 1
    assert non_markers[0]["study_design"] == "RCT"


def test_exclude_empty_omits_empty_papers(db_with_empty_extractions, spec, tmp_path):
    """With exclude_empty=True, papers with no extraction data are omitted."""
    out = str(tmp_path / "evidence.csv")
    export_evidence_csv(db_with_empty_extractions, spec, out, exclude_empty=True)

    with open(out) as f:
        reader = csv.DictReader(f)
        rows = list(reader)

    assert len(rows) == 1
    assert rows[0]["study_design"] == "RCT"


def test_exclude_empty_excel(db_with_empty_extractions, spec, tmp_path):
    """Excel export also supports exclude_empty."""
    out = str(tmp_path / "evidence.xlsx")
    export_evidence_excel(db_with_empty_extractions, spec, out, exclude_empty=True)

    wb = openpyxl.load_workbook(out)
    ws = wb["Evidence Table"]
    # 1 header + 1 data row (empty paper excluded)
    assert ws.max_row == 2
    wb.close()


# ── H15: Dynamic model names in methods section ─────────────────────


def test_methods_uses_spec_screening_model(populated_db, spec):
    """Methods section uses screening model from spec, not hardcoded."""
    methods = generate_methods_section(populated_db, spec)
    assert spec.screening_models.primary in methods


def test_methods_uses_db_extraction_model(populated_db, spec):
    """Methods section uses extraction model from DB data."""
    methods = generate_methods_section(populated_db, spec)
    assert "deepseek-r1:32b" in methods


def test_methods_uses_db_audit_model(populated_db, spec):
    """Methods section uses auditor model from DB evidence_spans."""
    methods = generate_methods_section(populated_db, spec)
    assert "qwen3:32b" in methods


def test_methods_multi_model_ft_screening(tmp_path, spec):
    """Methods section reports multiple FT screening models with counts."""
    db = ReviewDatabase("test_ft_models", data_root=tmp_path)

    cits = [
        Citation(title=f"Study {i}", source="pubmed", pmid=str(i),
                 doi=f"10.1/{i}", authors=["Auth A"], journal="J Test", year=2023)
        for i in range(1, 6)
    ]
    db.add_papers(cits)
    papers = db.get_papers_by_status("INGESTED")

    for p in papers:
        pid = p["id"]
        db.add_screening_decision(pid, 1, "include", "Relevant", "qwen3:8b")
        db.add_screening_decision(pid, 2, "include", "Confirmed", "qwen3:8b")
        db.update_status(pid, "ABSTRACT_SCREENED_IN")
        db.update_status(pid, "PDF_ACQUIRED")
        db.update_status(pid, "PARSED")

    # Add FT screening decisions with two different models
    from engine.core.database import _now
    for i, p in enumerate(papers):
        model = "qwen3.5:27b" if i < 3 else "qwen3:32b"
        db._conn.execute(
            """INSERT INTO ft_screening_decisions
               (paper_id, model, decision, reason_code, rationale, confidence, decided_at)
               VALUES (?, ?, ?, ?, ?, ?, ?)""",
            (p["id"], model, "FT_ELIGIBLE", "IN_SCOPE", "Relevant", 0.9, _now()),
        )
        db.update_status(p["id"], "FT_ELIGIBLE")

    schema_hash = spec.extraction_hash()
    for p in papers:
        pid = p["id"]
        db.update_status(pid, "EXTRACTED")
        ext_id = db.add_extraction(pid, schema_hash, {}, "trace", "deepseek-r1:32b")
        db.add_evidence_span(ext_id, "study_design", "RCT", ".", 0.9)
        spans = db._conn.execute(
            "SELECT id FROM evidence_spans WHERE extraction_id = ?", (ext_id,)
        ).fetchall()
        for s in spans:
            db.update_audit(s["id"], "verified", "gemma3:27b", "OK")
        db.update_status(pid, "AI_AUDIT_COMPLETE")

    db._conn.commit()

    methods = generate_methods_section(db, spec)

    # Should contain both FT models with counts
    assert "qwen3.5:27b (n=3)" in methods
    assert "qwen3:32b (n=2)" in methods

    db.close()


def test_methods_placeholder_when_no_data(tmp_path, spec):
    """Methods section uses [MODEL NOT SPECIFIED] when DB has no extraction data."""
    db = ReviewDatabase("test_placeholder", data_root=tmp_path)

    cits = [
        Citation(title="Study 1", source="pubmed", pmid="1",
                 doi="10.1/1", authors=["Auth A"], journal="J Test", year=2023)
    ]
    db.add_papers(cits)
    papers = db.get_papers_by_status("INGESTED")
    for p in papers:
        db.add_screening_decision(p["id"], 1, "include", "Relevant", "qwen3:8b")
        db.add_screening_decision(p["id"], 2, "include", "OK", "qwen3:8b")
        db.update_status(p["id"], "ABSTRACT_SCREENED_IN")
        db.update_status(p["id"], "PDF_ACQUIRED")
        db.update_status(p["id"], "PARSED")
        db.update_status(p["id"], "EXTRACTED")
        db.update_status(p["id"], "AI_AUDIT_COMPLETE")

    methods = generate_methods_section(db, spec)

    # No extractions or audits in DB → placeholders
    assert "[MODEL NOT SPECIFIED]" in methods

    db.close()
