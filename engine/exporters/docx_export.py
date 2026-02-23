"""DOCX formatted evidence table for journal submission."""

import json
import logging

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Inches, Pt

from engine.core.database import ReviewDatabase
from engine.core.review_spec import ReviewSpec

logger = logging.getLogger(__name__)


def export_evidence_docx(
    db: ReviewDatabase, spec: ReviewSpec, output_path: str
) -> None:
    """Export a professional evidence table as DOCX."""
    doc = Document()

    # Page setup: landscape, narrow margins
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    # Title
    title_para = doc.add_paragraph()
    run = title_para.add_run(spec.title)
    run.bold = True
    run.font.size = Pt(14)
    title_para.add_run(f"\nVersion {spec.version} — {spec.date}")

    doc.add_paragraph("")  # spacer

    # Build table columns: fixed columns + extraction field columns
    field_names = [f.name for f in spec.extraction_schema.fields]
    base_cols = ["Study", "Year", "Journal"]
    all_cols = base_cols + field_names

    # Create table
    papers = db._conn.execute(
        "SELECT * FROM papers WHERE status IN ('EXTRACTED', 'AUDITED') ORDER BY id"
    ).fetchall()

    table = doc.add_table(rows=1 + len(papers), cols=len(all_cols))
    table.style = "Table Grid"

    # Header row
    for i, col_name in enumerate(all_cols):
        cell = table.rows[0].cells[i]
        cell.text = col_name.replace("_", " ").title()
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)

    # Data rows
    for row_idx, paper in enumerate(papers, 1):
        pid = paper["id"]

        # Get latest extraction spans
        extraction = db._conn.execute(
            "SELECT id FROM extractions WHERE paper_id = ? ORDER BY id DESC LIMIT 1",
            (pid,),
        ).fetchone()

        span_map = {}
        if extraction:
            spans = db._conn.execute(
                "SELECT field_name, value FROM evidence_spans WHERE extraction_id = ?",
                (extraction["id"],),
            ).fetchall()
            for s in spans:
                span_map[s["field_name"]] = s["value"]

        # Authors: first author et al.
        authors_raw = paper["authors"] or "[]"
        try:
            authors = json.loads(authors_raw)
        except (json.JSONDecodeError, TypeError):
            authors = []
        if authors:
            study_label = f"{authors[0].split()[-1]} et al." if len(authors) > 1 else authors[0]
        else:
            study_label = paper["title"][:40]

        base_values = [study_label, str(paper["year"] or ""), paper["journal"] or ""]
        field_values = [span_map.get(f, "") for f in field_names]
        all_values = base_values + field_values

        for col_idx, val in enumerate(all_values):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = str(val) if val else ""
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    doc.save(output_path)
    logger.info("Evidence DOCX exported to %s (%d studies)", output_path, len(papers))
